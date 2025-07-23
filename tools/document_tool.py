"""Document tool for loading PRD content from Word documents."""

import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
from docx import Document
from docx.shared import Inches
import docx2txt
import io

logger = logging.getLogger(__name__)

class DocumentTool:
    """Tool for loading and processing Word documents containing PRD content."""
    
    def __init__(self):
        """Initialize the Document tool."""
        logger.info("DocumentTool initialized")
    
    def load_document_content(self, file_path: str) -> Dict[str, Any]:
        """
        Load content from a Word document (.doc or .docx).
        
        Args:
            file_path: Path to the Word document file
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            if not self.validate_document_file(file_path):
                raise ValueError(f"Invalid document file: {file_path}")
            
            logger.info(f"Loading Word document from: {file_path}")
            
            # Detect file format and process accordingly
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.docx':
                # Use python-docx for .docx files
                doc = Document(file_path)
                content_data = {
                    'title': self._extract_title(doc),
                    'content': self._extract_text_content(doc),
                    'sections': self._extract_sections(doc),
                    'tables': self._extract_tables(doc),
                    'metadata': self._extract_metadata(doc, file_path),
                    'summary': self._generate_content_summary(doc)
                }
            elif file_extension == '.doc':
                # Use docx2txt for .doc files
                text_content = docx2txt.process(file_path)
                content_data = {
                    'title': Path(file_path).stem,
                    'content': text_content,
                    'sections': self._extract_sections_from_text(text_content),
                    'tables': [],  # Limited table extraction for .doc files
                    'metadata': {'filename': Path(file_path).name, 'format': 'doc'},
                    'summary': self._generate_text_summary(text_content)
                }
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            logger.info(f"Successfully loaded document with {len(content_data['content'])} characters")
            return content_data
            
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            raise
    
    def load_document_from_bytes(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Load content from Word document bytes (for uploaded files).
        
        Args:
            file_bytes: Document file content as bytes
            filename: Original filename
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            logger.info(f"Loading Word document from bytes: {filename}")
            
            # Detect file format from filename
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.docx':
                # Use python-docx for .docx files
                doc = Document(io.BytesIO(file_bytes))
                content_data = {
                    'title': self._extract_title(doc) or filename,
                    'content': self._extract_text_content(doc),
                    'sections': self._extract_sections(doc),
                    'tables': self._extract_tables(doc),
                    'metadata': self._extract_metadata_from_bytes(doc, filename),
                    'summary': self._generate_content_summary(doc)
                }
            elif file_extension == '.doc':
                # For .doc files, we need to save bytes to temp file first
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                    temp_file.write(file_bytes)
                    temp_file_path = temp_file.name
                
                try:
                    text_content = docx2txt.process(temp_file_path)
                    content_data = {
                        'title': Path(filename).stem,
                        'content': text_content,
                        'sections': self._extract_sections_from_text(text_content),
                        'tables': [],  # Limited table extraction for .doc files
                        'metadata': {'filename': filename, 'format': 'doc'},
                        'summary': self._generate_text_summary(text_content)
                    }
                finally:
                    # Clean up temp file
                    import os
                    try:
                        os.unlink(temp_file_path)
                    except OSError:
                        pass
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            logger.info(f"Successfully loaded document with {len(content_data['content'])} characters")
            return content_data
            
        except Exception as e:
            logger.error(f"Error loading document from bytes: {e}")
            raise
    
    def _extract_title(self, doc: Document) -> str:
        """Extract document title."""
        try:
            # Try to get title from document properties
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                return doc.core_properties.title
            
            # Fallback: use first paragraph if it looks like a title
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                first_para = doc.paragraphs[0].text.strip()
                if len(first_para) < 100 and not first_para.endswith('.'):
                    return first_para
            
            return "Document Content"
            
        except Exception as e:
            logger.warning(f"Error extracting title: {e}")
            return "Document Content"
    
    def _extract_text_content(self, doc: Document) -> str:
        """Extract all text content from the document."""
        try:
            content_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            return "\n\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text content: {e}")
            return ""
    
    def _extract_sections(self, doc: Document) -> List[Dict[str, str]]:
        """Extract sections based on heading styles."""
        try:
            sections = []
            current_section = None
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Check if paragraph is a heading
                style_name = paragraph.style.name.lower()
                if 'heading' in style_name or paragraph.runs and paragraph.runs[0].bold:
                    # Save previous section
                    if current_section:
                        sections.append(current_section)
                    
                    # Start new section
                    current_section = {
                        'title': text,
                        'content': ''
                    }
                else:
                    # Add to current section
                    if current_section:
                        if current_section['content']:
                            current_section['content'] += "\n"
                        current_section['content'] += text
                    else:
                        # No section started yet, create a default one
                        if not sections:
                            current_section = {
                                'title': 'Introduction',
                                'content': text
                            }
            
            # Add last section
            if current_section:
                sections.append(current_section)
            
            return sections
            
        except Exception as e:
            logger.error(f"Error extracting sections: {e}")
            return []
    
    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract tables from the document."""
        try:
            tables_data = []
            
            for i, table in enumerate(doc.tables):
                table_data = {
                    'table_number': i + 1,
                    'rows': [],
                    'text_summary': ''
                }
                
                rows_text = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data['rows'].append(row_data)
                    rows_text.append(" | ".join(row_data))
                
                table_data['text_summary'] = "\n".join(rows_text)
                tables_data.append(table_data)
            
            return tables_data
            
        except Exception as e:
            logger.error(f"Error extracting tables: {e}")
            return []
    
    def _extract_metadata(self, doc: Document, file_path: str) -> Dict[str, Any]:
        """Extract document metadata."""
        try:
            file_path = Path(file_path)
            metadata = {
                'filename': file_path.name,
                'file_size': file_path.stat().st_size if file_path.exists() else 0,
                'file_type': 'Word Document',
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables)
            }
            
            # Add document properties if available
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata['author'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                metadata['created'] = doc.core_properties.created.isoformat()
            if hasattr(doc.core_properties, 'modified') and doc.core_properties.modified:
                metadata['modified'] = doc.core_properties.modified.isoformat()
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {e}")
            return {'filename': 'unknown', 'file_type': 'Word Document'}
    
    def _extract_metadata_from_bytes(self, doc: Document, filename: str) -> Dict[str, Any]:
        """Extract document metadata from bytes."""
        try:
            metadata = {
                'filename': filename,
                'file_type': 'Word Document',
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables)
            }
            
            # Add document properties if available
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata['author'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                metadata['created'] = doc.core_properties.created.isoformat()
            if hasattr(doc.core_properties, 'modified') and doc.core_properties.modified:
                metadata['modified'] = doc.core_properties.modified.isoformat()
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {e}")
            return {'filename': filename, 'file_type': 'Word Document'}
    
    def _generate_content_summary(self, doc: Document) -> str:
        """Generate a summary of the document content."""
        try:
            content = self._extract_text_content(doc)
            
            # Basic summary stats
            word_count = len(content.split())
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            summary_parts = [
                f"Document contains {word_count} words in {paragraph_count} paragraphs"
            ]
            
            if table_count > 0:
                summary_parts.append(f"and {table_count} table(s)")
            
            # Extract key topics (simple approach)
            sections = self._extract_sections(doc)
            if sections:
                section_titles = [section['title'] for section in sections[:5]]
                summary_parts.append(f"Main sections: {', '.join(section_titles)}")
            
            return ". ".join(summary_parts) + "."
            
        except Exception as e:
            logger.error(f"Error generating summary: {e}")
            return "Document summary not available."
    
    def validate_document_file(self, file_path: str) -> bool:
        """
        Validate if the file is a valid Word document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            True if valid, False otherwise
        """
        try:
            file_path = Path(file_path)
            
            # Check if file exists
            if not file_path.exists():
                logger.error(f"Document file does not exist: {file_path}")
                return False
            
            # Check file extension
            valid_extensions = ['.docx', '.doc']
            if file_path.suffix.lower() not in valid_extensions:
                logger.error(f"Invalid file extension: {file_path.suffix}")
                return False
            
            # Try to open the document based on format
            try:
                if file_path.suffix.lower() == '.docx':
                    Document(file_path)
                elif file_path.suffix.lower() == '.doc':
                    # For .doc files, just check if docx2txt can process it
                    docx2txt.process(str(file_path))
                return True
            except Exception as e:
                logger.error(f"Cannot open document: {e}")
                return False
                
        except Exception as e:
            logger.error(f"Error validating document file: {e}")
            return False
    
    def validate_document_bytes(self, file_bytes: bytes) -> bool:
        """
        Validate if the bytes represent a valid Word document.
        
        Args:
            file_bytes: Document content as bytes
            
        Returns:
            True if valid, False otherwise
        """
        try:
            Document(io.BytesIO(file_bytes))
            return True
        except Exception as e:
            logger.error(f"Cannot open document from bytes: {e}")
            return False
    
    def extract_text_for_analysis(self, content_data: Dict[str, Any]) -> str:
        """
        Extract and format text content suitable for LLM analysis.
        
        Args:
            content_data: Document content data from load_document_content
            
        Returns:
            Formatted text for analysis
        """
        try:
            analysis_parts = []
            
            # Add title
            if content_data.get('title'):
                analysis_parts.append(f"DOCUMENT TITLE: {content_data['title']}")
                analysis_parts.append("=" * 50)
            
            # Add main content
            if content_data.get('content'):
                analysis_parts.append("DOCUMENT CONTENT:")
                analysis_parts.append(content_data['content'])
                analysis_parts.append("")
            
            # Add sections if available
            sections = content_data.get('sections', [])
            if sections:
                analysis_parts.append("DOCUMENT SECTIONS:")
                for i, section in enumerate(sections, 1):
                    analysis_parts.append(f"\n{i}. {section['title']}")
                    if section.get('content'):
                        analysis_parts.append(section['content'])
                analysis_parts.append("")
            
            # Add tables if available
            tables = content_data.get('tables', [])
            if tables:
                analysis_parts.append("TABLES:")
                for table in tables:
                    analysis_parts.append(f"\nTable {table['table_number']}:")
                    analysis_parts.append(table['text_summary'])
                analysis_parts.append("")
            
            # Add metadata
            metadata = content_data.get('metadata', {})
            if metadata:
                analysis_parts.append("DOCUMENT METADATA:")
                for key, value in metadata.items():
                    analysis_parts.append(f"- {key}: {value}")
            
            return "\n".join(analysis_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text for analysis: {e}")
            return content_data.get('content', '')
    
    def _extract_sections_from_text(self, text_content: str) -> List[Dict[str, str]]:
        """
        Extract sections from plain text (for .doc files).
        
        Args:
            text_content: Plain text content
            
        Returns:
            List of section dictionaries
        """
        try:
            sections = []
            lines = text_content.split('\n')
            current_section = None
            section_content = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Simple heuristic: lines that are all caps or start with numbers might be headers
                if (line.isupper() and len(line.split()) <= 5) or \
                   (line[0].isdigit() and '.' in line[:10]):
                    # Save previous section
                    if current_section and section_content:
                        sections.append({
                            'heading': current_section,
                            'content': '\n'.join(section_content)
                        })
                    
                    # Start new section
                    current_section = line
                    section_content = []
                else:
                    section_content.append(line)
            
            # Add final section
            if current_section and section_content:
                sections.append({
                    'heading': current_section,
                    'content': '\n'.join(section_content)
                })
            
            return sections[:10]  # Limit to first 10 sections
            
        except Exception as e:
            logger.error(f"Error extracting sections from text: {e}")
            return []
    
    def _generate_text_summary(self, text_content: str) -> str:
        """
        Generate a summary from plain text content (for .doc files).
        
        Args:
            text_content: Plain text content
            
        Returns:
            Content summary
        """
        try:
            if not text_content:
                return "Empty document"
            
            lines = [line.strip() for line in text_content.split('\n') if line.strip()]
            
            # Basic statistics
            word_count = len(text_content.split())
            line_count = len(lines)
            
            # Try to find key sections
            summary_parts = []
            summary_parts.append(f"Document contains {word_count} words across {line_count} lines")
            
            # Look for potential headers or important lines
            important_lines = []
            for line in lines[:20]:  # Check first 20 lines
                if (line.isupper() and len(line.split()) <= 8) or \
                   any(keyword in line.lower() for keyword in ['overview', 'summary', 'introduction', 'purpose']):
                    important_lines.append(line)
            
            if important_lines:
                summary_parts.append(f"Key sections identified: {', '.join(important_lines[:3])}")
            
            return ". ".join(summary_parts) + "."
            
        except Exception as e:
            logger.error(f"Error generating text summary: {e}")
            return "Document summary not available." 